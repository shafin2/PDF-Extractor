import pdfplumber
import pandas as pd
from docx import Document

def extract_and_print_pdf_content_to_doc(pdf_path, doc_path):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            doc.add_paragraph(f"--- Page {page_num + 1} ---")
            words = page.extract_words()
            tables = page.extract_tables()
            table_bboxes = [table.bbox for table in page.find_tables()]
            content = []
            for word in words:
                if not any(is_word_in_bbox(word, bbox) for bbox in table_bboxes):
                    content.append({
                        "type": "text",
                        "text": word["text"],
                        "position": word["top"], 
                        "x0": word["x0"] 
                    })
            for i, table in enumerate(tables):
                df = pd.DataFrame(table)
                content.append({
                    "type": "table",
                    "table": df,
                    "position": table_bboxes[i][1], 
                })
            content.sort(key=lambda x: (x["position"], x["x0"] if "x0" in x else 0))
            previous_position = None
            line_text = []
            for item in content:
                if item["type"] == "text":
                    if previous_position is not None and item["position"] != previous_position:
                        doc.add_paragraph(" ".join(line_text)) 
                        line_text = [] 
                    line_text.append(item["text"])
                    previous_position = item["position"]

                elif item["type"] == "table":
                    if line_text:
                        doc.add_paragraph(" ".join(line_text))
                        line_text = []

                    doc.add_paragraph("\n--- Table ---")
                    word_table = doc.add_table(rows=len(item["table"]), cols=len(item["table"].columns))
                    for row_idx, row in enumerate(item["table"].values):
                        for col_idx, cell in enumerate(row):
                            word_table.cell(row_idx, col_idx).text = str(cell)
                    doc.add_paragraph()  

            if line_text:
                doc.add_paragraph(" ".join(line_text))


    doc.save(doc_path)

def is_word_in_bbox(word, bbox):
    x0, y0, x1, y1 = bbox
    word_x0, word_y0, word_x1, word_y1 = word["x0"], word["top"], word["x1"], word["bottom"]
    # Check if the word falls within the bounding box of a table
    return (word_x0 >= x0 and word_x1 <= x1 and word_y0 >= y0 and word_y1 <= y1)

pdf_path = "SamplePdfs\\Spring 24 - Platform.pdf"
doc_path = f"{pdf_path}-output.docx"

extract_and_print_pdf_content_to_doc(pdf_path, doc_path)